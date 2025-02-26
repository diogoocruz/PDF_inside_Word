import streamlit as st
import tempfile
import os
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from docx2pdf import convert


def add_pdf_images_to_docx(template_path, pdf_path, output_path, output_format="docx", dpi=200, img_width_inches=7):
    # Carrega o documento a partir do template
    doc = Document(template_path)
    
    # Se existir um parágrafo, utiliza o primeiro; caso contrário, cria um novo
    if doc.paragraphs:
        first_paragraph = doc.paragraphs[0]
    else:
        first_paragraph = doc.add_paragraph()
    
    # Converte as páginas do PDF para imagens
    images = convert_from_path(pdf_path, dpi=dpi)
    
    # Cria uma pasta temporária para salvar as imagens
    temp_dir = os.path.join(os.path.dirname(output_path), "temp_images")
    os.makedirs(temp_dir, exist_ok=True)
    
    # Para cada imagem, salva em arquivo temporário e insere no documento
    for idx, img in enumerate(images, start=1):
        temp_img_path = os.path.join(temp_dir, f"pagina_{idx}.png")
        img.save(temp_img_path, "PNG")
        
        if idx == 1:
            run = first_paragraph.add_run()
            run.add_picture(temp_img_path, width=Inches(img_width_inches))
        else:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(temp_img_path, width=Inches(img_width_inches))
    
    # Salva o documento como .docx
    doc.save(output_path)
    
    # Se solicitado, converte para PDF
    if output_format.lower() == "pdf":
        # Cria o caminho para o PDF de saída
        pdf_output = output_path.replace(".docx", ".pdf")
        convert(output_path, pdf_output)
    
    # Remove os arquivos temporários
    for filename in os.listdir(temp_dir):
        os.remove(os.path.join(temp_dir, filename))
    os.rmdir(temp_dir)

# Interface Streamlit
st.set_page_config(page_title="Inserir PDFs em Templates Word", page_icon="📄")

st.title("Incorpore PDFs em Templates Word Personalizados")
st.write("Insira os seus PDFs dentro de um template em DOCX.")
st.subheader("Instruções")
st.write("Faça upload do template (DOCX) e do arquivo PDF. Configure as opções e clique em Converter.")

uploaded_template = st.file_uploader("Template DOCX", type=["docx"])
uploaded_pdf = st.file_uploader("Arquivo PDF", type=["pdf"])

output_format = st.selectbox("Formato de saída", ["docx", "pdf"])
dpi = st.number_input("DPI (qualidade da imagem)", min_value=50, max_value=600, value=200, step=10)
img_width = st.number_input("Largura da imagem (polegadas)", min_value=1.0, max_value=20.0, value=7.0, step=0.5)

if st.button("Converter"):
    if not uploaded_template or not uploaded_pdf:
        st.error("Por favor, faça upload dos dois arquivos!")
    else:
        with tempfile.TemporaryDirectory() as tmpdirname:
            # Define caminhos temporários para os arquivos
            template_path = os.path.join(tmpdirname, "template.docx")
            pdf_path = os.path.join(tmpdirname, "input.pdf")
            output_docx_path = os.path.join(tmpdirname, "output.docx")
            
            # Salva os arquivos enviados
            with open(template_path, "wb") as f:
                f.write(uploaded_template.read())
            with open(pdf_path, "wb") as f:
                f.write(uploaded_pdf.read())
            
            # Realiza a conversão
            add_pdf_images_to_docx(template_path, pdf_path, output_docx_path,
                                   output_format=output_format, dpi=dpi, img_width_inches=img_width)
            
            # Define o caminho do arquivo final
            if output_format.lower() == "pdf":
                final_output = output_docx_path.replace(".docx", ".pdf")
            else:
                final_output = output_docx_path
            
            # Lê o arquivo final e disponibiliza para download
            with open(final_output, "rb") as f:
                st.download_button("Baixar Arquivo Gerado",
                                   data=f.read(),
                                   file_name=os.path.basename(final_output))
